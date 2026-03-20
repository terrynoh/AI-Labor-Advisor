# -*- coding: utf-8 -*-
"""
pdf_generator.py
ระบบสร้าง PDF สำหรับ:
  1. แบบฟอร์ม คร.7 (คำร้องต่อพนักงานตรวจแรงงาน)
  2. หนังสือบอกกล่าวทวงถาม (Demand Letter)
"""

import logging
import os
import platform
import shutil
import subprocess
import tempfile
from datetime import datetime

logger = logging.getLogger(__name__)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── LibreOffice command path ────────────────────────────────────────────────
if platform.system() == "Windows":
    LIBREOFFICE_CMD = r"C:\Program Files\LibreOffice\program\soffice.exe"
else:
    LIBREOFFICE_CMD = "libreoffice"

# ── ที่อยู่ template ──────────────────────────────────────────────────────
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template_kor7.docx")

# ── ตัวเลขไทย ──────────────────────────────────────────────────────────────
THAI_MONTHS = [
    "", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
    "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
]

def _to_thai_year(year: int) -> int:
    return year + 543

def _today_parts():
    now = datetime.now()
    return str(now.day), THAI_MONTHS[now.month], str(_to_thai_year(now.year))

def _baht_text(amount) -> str:
    try:
        n = float(amount)
        return f"{n:,.2f} บาท"
    except (ValueError, TypeError):
        return str(amount)

def _parse_date(date_str) -> str:
    """
    여러 날짜 포맷 → 'DD เดือน YYYY(พ.ศ.)' 형식으로 변환
    지원 포맷:
      DD/MM/YYYY  (웹 폼 → fmtDate 변환 결과)
      DDMMYYYY    (숫자만)
      YYYY-MM-DD  (HTML date input 직접값 혹시 올 경우)
    연도는 서기/불기 모두 허용 (2000 미만이면 불기로 간주)
    """
    if not date_str:
        return ""
    s = str(date_str).strip()

    # YYYY-MM-DD
    if len(s) == 10 and s[4] == "-":
        parts = s.split("-")
        if len(parts) == 3:
            dd, mm, yyyy = int(parts[2]), int(parts[1]), parts[0]
            if 1 <= mm <= 12:
                return f"{dd} {THAI_MONTHS[mm]} {yyyy}"

    # DD/MM/YYYY 또는 구분자 제거 후 DDMMYYYY
    digits = s.replace("/", "").replace("-", "").replace(" ", "")
    if len(digits) == 8 and digits.isdigit():
        dd   = int(digits[0:2])
        mm   = int(digits[2:4])
        yyyy = digits[4:8]
        if 1 <= mm <= 12:
            return f"{dd} {THAI_MONTHS[mm]} {yyyy}"

    return s  # 파싱 불가 시 원본 반환

def _fill(run, value):
    if value:
        run.text = f" {str(value)} "
    else:
        run.text = "................................"

def _libreoffice_convert(tmp_dir, tmp_docx, out_pdf_name):
    """Convert docx to PDF via LibreOffice, return PDF path."""
    result = subprocess.run(
        [LIBREOFFICE_CMD, "--headless", "--convert-to", "pdf",
         "--outdir", tmp_dir, tmp_docx],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr}")
    tmp_pdf = os.path.join(tmp_dir, out_pdf_name)
    if not os.path.exists(tmp_pdf):
        raise RuntimeError("PDF conversion failed — output file not found")
    return tmp_pdf


# ═══════════════════════════════════════════════════════════════════════════
#  1. คร.7 진정서
# ═══════════════════════════════════════════════════════════════════════════

def generate_kor7_pdf(data: dict, output_path: str) -> str:
    doc = Document(TEMPLATE_PATH)
    paras = doc.paragraphs
    d = data

    def f(para_idx, run_idx, value):
        try:
            _fill(paras[para_idx].runs[run_idx], value)
        except (IndexError, AttributeError):
            pass

    day, month, year = _today_parts()
    f(8,  2, d.get("filed_province", "กรุงเทพมหานคร"))
    f(9,  1, day)
    f(9,  3, month)
    f(9,  7, year)

    f(10, 7, d.get("complainant_name"))
    f(11, 1, d.get("id_number"))
    f(12, 1, d.get("id_issued_at"))
    f(12, 3, d.get("id_issue_date"))
    f(12, 5, d.get("id_expiry"))
    f(13, 1, d.get("nationality", "ไทย"))
    f(13, 5, d.get("work_permit"))
    f(14, 1, d.get("age"))
    f(14, 3, d.get("address_no"))
    f(14, 5, d.get("moo"))
    f(14, 7, d.get("street"))
    f(15, 3, d.get("subdistrict"))
    f(15, 7, d.get("district"))
    f(16, 1, d.get("province"))
    f(16, 3, d.get("postal_code"))
    f(16, 5, d.get("phone"))

    f(17, 1, d.get("reg_address_no", d.get("address_no")))
    f(17, 3, d.get("reg_moo",        d.get("moo")))
    f(17, 5, d.get("reg_street",     d.get("street")))
    f(18, 3, d.get("reg_subdistrict",d.get("subdistrict")))
    f(18, 7, d.get("reg_district",   d.get("district")))
    f(19, 1, d.get("reg_province",   d.get("province")))
    f(19, 3, d.get("reg_postal_code",d.get("postal_code")))
    f(19, 5, d.get("reg_phone",      d.get("phone")))

    f(27, 4, d.get("employer_name"))
    f(28, 7, d.get("employer_person"))
    f(28, 9, d.get("business_type"))
    f(29, 1, d.get("employer_address_no"))
    f(29, 3, d.get("employer_moo"))
    f(29, 5, d.get("employer_street"))
    f(29, 9, d.get("employer_subdistrict"))
    f(30, 3, d.get("employer_district"))
    f(30, 5, d.get("employer_province"))
    f(30, 7, d.get("employer_postal"))
    f(31, 1, d.get("employer_phone"))
    f(31, 3, d.get("employer_landmark"))

    f(36, 1, d.get("workplace_address_no",  d.get("employer_address_no")))
    f(36, 3, d.get("workplace_moo",         d.get("employer_moo")))
    f(36, 5, d.get("workplace_street",      d.get("employer_street")))
    f(36, 9, d.get("workplace_subdistrict", d.get("employer_subdistrict")))
    f(37, 3, d.get("workplace_district",    d.get("employer_district")))
    f(37, 5, d.get("workplace_province",    d.get("employer_province")))
    f(37, 7, d.get("workplace_postal",      d.get("employer_postal")))
    f(38, 1, d.get("workplace_phone",       d.get("employer_phone")))
    f(38, 3, d.get("workplace_landmark",    d.get("employer_landmark")))

    # 날짜 파싱 적용
    start_date = _parse_date(d.get("start_date"))
    end_date   = _parse_date(d.get("end_date"))

    f(39, 4, start_date)
    f(39, 6, end_date)
    f(40, 1, d.get("position"))
    f(40, 5, d.get("department"))
    f(41, 5, d.get("supervisor"))
    f(41, 7, d.get("wage_rate"))

    f(44, 4, d.get("work_days_per_week"))
    f(44, 6, d.get("work_hours_per_day"))
    f(44, 8, d.get("work_start_time"))
    f(45, 1, d.get("work_end_time"))
    f(45, 5, d.get("break_start_time"))
    f(45, 9, d.get("break_end_time"))

    f(47, 6, d.get("pay_schedule_wage"))
    f(48, 6, d.get("pay_schedule_overtime"))
    f(49, 6, d.get("pay_schedule_holiday"))
    f(50, 6, d.get("pay_schedule_holiday_ot"))

    reason = d.get("reason", "")
    if reason:
        words = reason.split()
        mid = len(words) // 2
        f(51, 4, " ".join(words[:mid]))
        f(52, 0, " ".join(words[mid:]))

    wage_owed = d.get("wage_owed")
    if wage_owed:
        f(54, 4, d.get("wage_from_date", d.get("start_date")))
        f(54, 6, d.get("wage_to_date",   d.get("end_date")))
        f(55, 1, f"{float(wage_owed):,.2f}")
        f(55, 3, _baht_text(wage_owed))

    min_wage = d.get("minimum_wage")
    if min_wage:
        f(56, 4, f"{float(min_wage):,.2f}")

    notice = d.get("notice_pay")
    if notice:
        f(58, 4, f"{float(notice):,.2f}")
        f(59, 0, _baht_text(notice))

    ot = d.get("ot_amount")
    if ot:
        f(60, 4, d.get("ot_from_date", d.get("start_date")))
        f(60, 6, d.get("ot_to_date",   d.get("end_date")))
        f(61, 1, f"{float(ot):,.2f}")
        f(61, 3, _baht_text(ot))

    hol = d.get("holiday_amount")
    if hol:
        f(62, 5, d.get("holiday_from_date", d.get("start_date")))
        f(62, 7, d.get("holiday_to_date",   d.get("end_date")))
        f(63, 1, f"{float(hol):,.2f}")
        f(63, 3, _baht_text(hol))

    hol_ot = d.get("holiday_ot_amount")
    if hol_ot:
        f(64, 3, d.get("holiday_ot_from_date", d.get("start_date")))
        f(64, 5, d.get("holiday_ot_to_date",   d.get("end_date")))
        f(65, 1, f"{float(hol_ot):,.2f}")
        f(65, 3, _baht_text(hol_ot))

    sev = d.get("severance")
    if sev:
        f(67, 4, _parse_date(d.get("severance_from_date", d.get("start_date"))))
        f(68, 1, _parse_date(d.get("severance_to_date",   d.get("end_date"))))
        f(68, 3, f"{float(sev):,.2f}")
        f(68, 5, _baht_text(sev))

    f(85, 3, d.get("filed_province", "กรุงเทพมหานคร"))
    f(91, 2, d.get("complainant_name"))
    try:
        paras[92].runs[0].text = f'\t\t\t\t\t\t\t         ({d.get("complainant_name", "")})'
    except (IndexError, AttributeError):
        pass

    tmp_dir  = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, "filled_kor7.docx")
    doc.save(tmp_docx)

    try:
        tmp_pdf = _libreoffice_convert(tmp_dir, tmp_docx, "filled_kor7.pdf")
    except FileNotFoundError:
        raise RuntimeError(f"LibreOffice not found at: {LIBREOFFICE_CMD}")
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice PDF 변환 타임아웃 (kor7)")
        raise RuntimeError("PDF 변환 시간 초과. 잠시 후 다시 시도해주세요.")
    except Exception as e:
        logger.error("kor7 PDF 변환 오류: %s", e, exc_info=True)
        raise RuntimeError("PDF 변환 중 오류가 발생했습니다.")

    shutil.move(tmp_pdf, output_path)
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
#  2. หนังสือบอกกล่าวทวงถาม (Demand Letter)
# ═══════════════════════════════════════════════════════════════════════════

def generate_demand_letter_pdf(data: dict, letter_body: str, output_path: str) -> str:
    """
    Generate a formal demand letter as PDF.

    data keys:
        complainant_name, age, address, phone,
        employer_name, employer_address,
        position, start_date, end_date, wage_rate,
        deadline (int, default 15),
        total_amount (float)

    letter_body: AI-generated body text (Thai)
    """
    now = datetime.now()
    thai_date = f"{now.day} {THAI_MONTHS[now.month]} {_to_thai_year(now.year)}"

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    def add_para(text="", bold=False, size=14, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold      = bold
            run.font.size = Pt(size)
            run.font.name = "TH Sarabun New"
        return p

    # Header
    add_para("หนังสือบอกกล่าวทวงถามและเรียกร้องสิทธิ์",
             bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("(ส่งทางไปรษณีย์ลงทะเบียนตอบรับ)",
             size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(f"วันที่ {thai_date}", size=14,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

    add_para("เรื่อง   ขอเรียกร้องสิทธิ์ตามกฎหมายคุ้มครองแรงงาน",
             bold=True, size=14, space_after=4)
    add_para(f"เรียน   {data.get('employer_name', '...')}", size=14, space_after=2)

    employer_addr = data.get("employer_address", "")
    if employer_addr:
        add_para(f"        {employer_addr}", size=14, space_after=12)

    # Opening
    opening = (
        f"ข้าพเจ้า {data.get('complainant_name', '...')} "
        f"อายุ {data.get('age', '...')} ปี "
        f"อยู่บ้านเลขที่ {data.get('address', '...')} "
        f"โทรศัพท์ {data.get('phone', '...')} "
        f"เคยเป็นลูกจ้างของ{data.get('employer_name', '...')} "
        f"ตำแหน่ง {data.get('position', '...')} "
        f"ระหว่างวันที่ {_parse_date(data.get('start_date', '...'))} "
        f"ถึง {_parse_date(data.get('end_date', '...'))} "
        f"ได้รับค่าจ้างอัตรา {data.get('wage_rate', '...')} บาทต่อเดือน"
    )
    add_para(opening, size=14, space_after=8)

    # AI body
    for paragraph in letter_body.split("\n"):
        if paragraph.strip():
            add_para(paragraph.strip(), size=14, space_after=6)

    # Demand
    deadline  = data.get("deadline", 15)
    total     = data.get("total_amount", 0)
    total_str = f"{float(total):,.2f}" if total else "..."

    add_para("", space_after=4)
    add_para(
        f"ข้าพเจ้าขอให้ท่านชำระเงินจำนวนรวม {total_str} บาท "
        f"ภายใน {deadline} วัน นับแต่วันที่ได้รับหนังสือฉบับนี้",
        bold=True, size=14, space_after=6
    )
    add_para("หากท่านเพิกเฉยหรือไม่ดำเนินการ ข้าพเจ้าจะดำเนินการ ดังนี้",
             size=14, space_after=4)
    add_para("1. ยื่นคำร้องต่อพนักงานตรวจแรงงาน กรมสวัสดิการและคุ้มครองแรงงาน",
             size=14, space_after=4)
    add_para("2. ฟ้องร้องดำเนินคดีทางแพ่งและอาญาต่อไป", size=14, space_after=12)
    add_para("จึงเรียนมาเพื่อทราบและดำเนินการโดยด่วน", size=14, space_after=16)

    # Signature
    add_para("ขอแสดงความนับถือ", size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para("ลงชื่อ ........................................", size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(f"      ({data.get('complainant_name', '...')})", size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para("      ผู้ร้อง", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    tmp_dir  = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, "demand_letter.docx")
    doc.save(tmp_docx)

    try:
        tmp_pdf = _libreoffice_convert(tmp_dir, tmp_docx, "demand_letter.pdf")
    except FileNotFoundError:
        raise RuntimeError(f"LibreOffice not found at: {LIBREOFFICE_CMD}")
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice PDF 변환 타임아웃 (demand_letter)")
        raise RuntimeError("PDF 변환 시간 초과. 잠시 후 다시 시도해주세요.")
    except Exception as e:
        logger.error("demand_letter PDF 변환 오류: %s", e, exc_info=True)
        raise RuntimeError("PDF 변환 중 오류가 발생했습니다.")

    shutil.move(tmp_pdf, output_path)
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return output_path
